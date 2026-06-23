"""Document → text segments with format-specific citation labels:
  PDF   → "s. 5"
  Word  → "sekcja: <nagłówek>" / "tabela 2"
  Excel → "arkusz Sprzedaż, w. 10–60"
  CSV   → "wiersze 10–60"
"""

import csv as csv_module
from dataclasses import dataclass
from io import BytesIO, StringIO

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader


@dataclass
class ParsedSegment:
    text: str
    page: int  # 1-based ordinal; stable sort key
    location: str  # human-readable citation label, used verbatim in answers


SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".xlsx", ".csv")
LEGACY_EXTENSIONS = (".doc", ".xls")  # known-but-unsupported (old binary Office formats)

_PDF_MAGIC = b"%PDF-"
_ZIP_MAGIC = b"PK\x03\x04"  # OOXML (.docx/.xlsx) are ZIP archives
_ROWS_PER_SEGMENT = 50  # rows per segment — keeps row-range citations accurate


def _ext(filename: str) -> str:
    name = filename.lower()
    dot = name.rfind(".")
    return name[dot:] if dot != -1 else ""


def validate_upload(filename: str, data: bytes) -> str | None:
    """Return a Polish error message if the upload is unacceptable, else None.

    Centralised here so the format allowlist lives next to the parsers it gates.
    """
    ext = _ext(filename)
    if ext in LEGACY_EXTENSIONS:
        return f"Format {ext} nie jest obsługiwany. Zapisz jako .docx/.xlsx i spróbuj ponownie."
    if ext not in SUPPORTED_EXTENSIONS:
        return "Obsługiwane formaty: PDF, Word (.docx), Excel (.xlsx), CSV."
    if ext == ".pdf" and not data.startswith(_PDF_MAGIC):
        return "Plik nie jest prawidłowym PDF-em."
    if ext in (".docx", ".xlsx") and not data.startswith(_ZIP_MAGIC):
        return f"Plik nie jest prawidłowym {ext} (uszkodzony lub w starym formacie binarnym)."
    return None  # CSV has no reliable signature — the parser validates it instead


def parse_pdf(data: bytes) -> list[ParsedSegment]:
    """Per-page text via pypdf."""
    reader = PdfReader(BytesIO(data))
    return [
        ParsedSegment(text=(p.extract_text() or "").strip(), page=i, location=f"s. {i}")
        for i, p in enumerate(reader.pages, start=1)
    ]


def parse_docx(data: bytes) -> list[ParsedSegment]:
    """Split on heading paragraphs; tables become their own segments."""
    doc = DocxDocument(BytesIO(data))
    segments: list[ParsedSegment] = []
    section_no = 0
    heading = "wprowadzenie"
    buffer: list[str] = []

    def flush() -> None:
        nonlocal section_no, buffer
        text = "\n".join(buffer).strip()
        if text:
            section_no += 1
            segments.append(
                ParsedSegment(text=text, page=section_no, location=f"sekcja: {heading}")
            )
        buffer = []

    for para in doc.paragraphs:
        txt = para.text.strip()
        if not txt:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading") or style.startswith("title") or style.startswith("nagłówek"):
            flush()  # close out the previous section before switching heading
            heading = txt[:60]
        else:
            buffer.append(txt)
    flush()

    for ti, table in enumerate(doc.tables, start=1):
        rows = [
            " | ".join(c.text.strip() for c in row.cells)
            for row in table.rows
            if any(c.text.strip() for c in row.cells)
        ]
        if rows:
            section_no += 1
            segments.append(
                ParsedSegment(text="\n".join(rows), page=section_no, location=f"tabela {ti}")
            )
    return segments


def _render_rows(header: list[str], rows: list[list[str]]) -> str:
    """Render rows as "Kolumna: wartość | …" lines — keeps column semantics for retrieval."""
    lines: list[str] = []
    for row in rows:
        pairs = [f"{h}: {v}" for h, v in zip(header, row, strict=False) if str(v).strip()]
        if pairs:
            lines.append(" | ".join(pairs))
    return "\n".join(lines).strip()


def parse_xlsx(data: bytes) -> list[ParsedSegment]:
    """Excel → one segment per sheet per row-window; location = "arkusz <name>, w. a–b"."""
    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    segments: list[ParsedSegment] = []
    ordinal = 0
    try:
        for ws in wb.worksheets:
            rows = [
                [("" if c is None else str(c)) for c in r] for r in ws.iter_rows(values_only=True)
            ]
            if not rows:
                continue
            header, body = rows[0], rows[1:]
            for start in range(0, len(body), _ROWS_PER_SEGMENT):
                window = body[start : start + _ROWS_PER_SEGMENT]
                text = _render_rows(header, window)
                if not text:
                    continue
                ordinal += 1
                first, last = start + 2, start + 1 + len(window)  # +1 header, 1-based rows
                loc = f"arkusz {ws.title}, w. {first}–{last}"
                segments.append(ParsedSegment(text=text, page=ordinal, location=loc))
    finally:
        wb.close()
    return segments


def parse_csv(data: bytes) -> list[ParsedSegment]:
    """CSV → one segment per row-window; delimiter sniffed (comma/semicolon/tab)."""
    text_data = data.decode("utf-8-sig", errors="replace")
    try:
        dialect: type[csv_module.Dialect] | csv_module.Dialect = csv_module.Sniffer().sniff(
            text_data[:2048], delimiters=",;\t"
        )
    except csv_module.Error:
        dialect = csv_module.excel
    rows = [[str(c) for c in row] for row in csv_module.reader(StringIO(text_data), dialect)]
    if not rows:
        return []

    header, body = rows[0], rows[1:]
    segments: list[ParsedSegment] = []
    ordinal = 0
    for start in range(0, len(body), _ROWS_PER_SEGMENT):
        window = body[start : start + _ROWS_PER_SEGMENT]
        text = _render_rows(header, window)
        if not text:
            continue
        ordinal += 1
        first, last = start + 2, start + 1 + len(window)
        segments.append(ParsedSegment(text=text, page=ordinal, location=f"wiersze {first}–{last}"))
    return segments


def parse_document(filename: str, data: bytes) -> list[ParsedSegment]:
    """Dispatch by extension. validate_upload must be called first."""
    ext = _ext(filename)
    if ext == ".pdf":
        return parse_pdf(data)
    if ext == ".docx":
        return parse_docx(data)
    if ext == ".xlsx":
        return parse_xlsx(data)
    if ext == ".csv":
        return parse_csv(data)
    raise ValueError(f"Unsupported file type: {ext or filename!r}")
